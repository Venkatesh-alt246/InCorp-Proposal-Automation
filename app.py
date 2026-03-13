from flask import Flask, request, send_file, jsonify
from flask_cors import CORS
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak, Image as RLImage, Flowable
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT, TA_JUSTIFY
from reportlab.pdfgen import canvas as pdfcanvas
from datetime import datetime
from pypdf import PdfWriter, PdfReader
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from reportlab.platypus import KeepTogether
import io
import os
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfbase.pdfmetrics import registerFontFamily
BASE_DIR = os.path.dirname(os.path.abspath(__file__))

app = Flask(__name__)
CORS(app)
application = app 

try:
    pdfmetrics.registerFont(TTFont('MicrosoftSansSerif', os.path.join(BASE_DIR, 'microsoftsansserif.ttf')))
    print("✅ Microsoft Sans Serif font registered successfully")
except Exception as e:
    print(f"⚠️ Font registration failed: {e}")

try:
    # micross.ttf is the italic variant of Microsoft Sans Serif
    pdfmetrics.registerFont(TTFont('MicrosoftSansSerif-Italic', os.path.join(BASE_DIR, 'micross.ttf')))
    print("✅ Microsoft Sans Serif Italic registered")
except Exception as e:
    try:
        # fallback: use regular as italic (no visual italic, but no crash)
        pdfmetrics.registerFont(TTFont('MicrosoftSansSerif-Italic', os.path.join(BASE_DIR, 'microsoftsansserif.ttf')))
        print("⚠️ Using regular as italic fallback for MicrosoftSansSerif")
    except:
        pass

try:
    pdfmetrics.registerFont(TTFont('Roboto', os.path.join(BASE_DIR, 'Roboto-Regular.ttf')))
    print("✅ Roboto Regular registered")
except Exception as e:
    print(f"⚠️ Font registration failed: {e}")

try:
    pdfmetrics.registerFont(TTFont('Roboto-Bold', os.path.join(BASE_DIR, 'Roboto-Bold.ttf')))
    print("✅ Roboto Bold registered")
except Exception as e:
    print(f"⚠️ Font registration failed: {e}")

try:
    pdfmetrics.registerFont(TTFont('Roboto-Italic', os.path.join(BASE_DIR, 'Roboto-Italic.ttf')))
    print("✅ Roboto Italic registered")
except Exception as e:
    print(f"⚠️ Roboto-Italic.ttf not found in {BASE_DIR} - italic text will use Roboto-Regular as fallback")
    print(f"   Download from: https://fonts.google.com/specimen/Roboto and place Roboto-Italic.ttf in the app folder")
    try:
        pdfmetrics.registerFont(TTFont('Roboto-Italic', os.path.join(BASE_DIR, 'Roboto-Regular.ttf')))
    except:
        pass

# Register font families so <b>, <i>, <b><i> tags work correctly in Paragraphs
try:
    registerFontFamily('Roboto',
        normal='Roboto',
        bold='Roboto-Bold',
        italic='Roboto-Italic',
        boldItalic='Roboto-Bold')
    print("✅ Roboto font family registered")
except Exception as e:
    print(f"⚠️ Roboto font family registration failed: {e}")

try:
    registerFontFamily('MicrosoftSansSerif',
        normal='MicrosoftSansSerif',
        bold='MicrosoftSansSerif',
        italic='MicrosoftSansSerif-Italic',
        boldItalic='MicrosoftSansSerif-Italic')
    print("✅ MicrosoftSansSerif font family registered")
except Exception as e:
    print(f"⚠️ MicrosoftSansSerif font family registration failed: {e}")
def format_currency(amount):
    """Format number as currency WITHOUT dollar sign - integers only"""
    if not amount or amount == '' or amount == '0':
        return ''
    try:
        return f'{int(float(amount)):,}'
    except:
        return ''

class CoverPageWithCompany(Flowable):
    """Custom flowable to draw cover image with company name overlay"""
    
    def __init__(self, image_path, company_name):
        Flowable.__init__(self)
        self.image_path = image_path
        self.company_name = company_name
    
    def wrap(self, availWidth, availHeight):
        """Tell ReportLab how much space this flowable needs"""
        return (availWidth, availHeight)
    
    def draw(self):
        c = self.canv
        page_width = letter[0]
        page_height = letter[1]
        
        if os.path.exists(self.image_path):
            c.saveState()
            c.translate(-0.5*inch, -1*inch)
            c.drawImage(self.image_path, 
                       0, 0,
                       width=page_width, 
                       height=page_height, 
                       preserveAspectRatio=True, 
                       mask='auto')
            c.restoreState()
        
        c.setFillColor(colors.white)
        c.setFont("MicrosoftSansSerif", 20)
        
        page_width = letter[0]  # Use actual page width
        text_width = c.stringWidth(self.company_name, "MicrosoftSansSerif", 20)
        x_position = (page_width - text_width) / 2  # Center on full page
        y_position = 1*inch  # Match the red banner position
        
        c.drawString(x_position, y_position, self.company_name)

class InCorpCanvas(pdfcanvas.Canvas):
    """Custom canvas with InCorp header and footer - SKIP PAGE 1"""
    
    def __init__(self, *args, **kwargs):
        pdfcanvas.Canvas.__init__(self, *args, **kwargs)
        self._saved_page_states = []

    def showPage(self):
        self._saved_page_states.append(dict(self.__dict__))
        self._startPage()

    def save(self):
        num_pages = len(self._saved_page_states)
        for i, state in enumerate(self._saved_page_states, start=1):
            self.__dict__.update(state)
            self.draw_header_footer(num_pages, i)
            pdfcanvas.Canvas.showPage(self)
        pdfcanvas.Canvas.save(self)

    def draw_header_footer(self, page_count, page_num):
        """Draw header and footer - SKIP FIRST PAGE COMPLETELY"""
       # page_num = len(self._saved_page_states)
        
        # CRITICAL: Skip page 1 entirely - no header, no footer, nothing
        if page_num == 1:
            return
        
        # For pages 2+, draw larger header
        try:
            header_path = os.path.join(BASE_DIR, 'incorp_header.png')
            if os.path.exists(header_path):
                self.drawImage(header_path, 
                             0.5 , letter[1] - 1*inch,
                             width=7.5*inch, height=1*inch,
                             preserveAspectRatio=False, mask='auto')
            else:
                self.setFillColor(colors.HexColor('#C00000'))
                self.rect(0.5*inch, letter[1] - 0.65*inch, 1.3*inch, 0.4*inch, fill=1, stroke=0)
                
                self.setFillColor(colors.HexColor('#44546A'))
                self.rect(1.75*inch, letter[1] - 0.65*inch, 1.3*inch, 0.4*inch, fill=1, stroke=0)
                
                self.setFillColor(colors.HexColor('#F5F5F5'))
                self.rect(3*inch, letter[1] - 0.65*inch, 5*inch, 0.4*inch, fill=1, stroke=0)
                
                self.setFillColor(colors.HexColor('#333333'))
                self.setFont("Helvetica-Bold", 14)
                self.drawRightString(letter[0]- 0.7*inch, letter[1] - 0.45*inch, "In.Corp")
                
                self.setFont("Helvetica", 6)
                self.setFillColor(colors.HexColor('#666666'))
                self.drawRightString(letter[0] - 0.7*inch, letter[1] - 0.55*inch, "An Ascentium Company")
        except Exception as e:
            print(f"Header warning: {e}")

        
        self.setStrokeColor(colors.HexColor('#CCCCCC'))
        self.setLineWidth(0.5)
        self.line(0.5*inch, 0.65*inch, letter[0] -0.5*inch, 0.65*inch)
        
        # Footer text - Left side
        self.setFillColor(colors.HexColor('#C00000'))
        self.setFont("Helvetica", 8)
        self.drawString(0.5*inch, 0.5*inch, "www.incorp.asia")
        
        adjusted_page_num = page_num + 3  # ✅ ADD 3 for static pages
        adjusted_total = page_count + 3 +9 
        # Footer text - Right side (page number)
        self.setFillColor(colors.HexColor('#C00000'))  # ✅ Add this line
        self.setFont("Helvetica", 8)
        self.drawRightString(2.2*inch, 0.5*inch, f"Page {adjusted_page_num} of {adjusted_total}")
        
        # Second line - copyright
        self.setFont("Helvetica", 7)
        self.setFillColor(colors.HexColor('#666666'))
        self.drawString(0.5*inch, 0.38*inch, "© In.Corp Global Pte Ltd. All Right Reserved.")
        
        # Third line - confidential notice
        self.setFillColor(colors.HexColor('#999999'))
        self.setFont("Helvetica", 7)
        self.drawString(0.5*inch, 0.26*inch, 
                      "This document is being furnished to you on a confidential basis and solely for your information.")

@app.route('/')
def index():
    """Serve the HTML form"""
    from flask import make_response
    response = make_response(send_file(os.path.join(BASE_DIR, 'incorp_form.html')))
    response.headers['Cache-Control'] = 'no-store, no-cache, must-revalidate, max-age=0'
    response.headers['Pragma'] = 'no-cache'
    response.headers['Expires'] = '0'
    return response


# ===================== FIX 1: build_pdf_elements is now its own TOP-LEVEL function =====================
def build_pdf_elements(data):
    """Generate dynamic PDF pages (1, 5-13) and return elements list"""
    # NO try/except here - just build and return elements
        
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        rightMargin=0.5*inch,
        leftMargin=0.5*inch,
        topMargin=1.2*inch,
        bottomMargin=1*inch
    )
    
    elements = []
    styles = getSampleStyleSheet()
    
    # SCOPE OF SERVICES / FEES main headings — large red bold
    heading1_style = ParagraphStyle(
        'Heading1Custom',
        parent=styles['Heading1'],
        fontSize=14,
        textColor=colors.HexColor('#BB2121'),
        fontName='Roboto-Bold',
        spaceAfter=10,
        spaceBefore=14,
        leading=18,
    )
    
    # A. B. C. D. sub-section headings — smaller red bold with indent
    heading2_style = ParagraphStyle(
        'Heading2Custom',
        parent=styles['Heading2'],
        fontSize=10,
        textColor=colors.HexColor('#C00000'),
        fontName='Roboto-Bold',
        spaceAfter=6,
        spaceBefore=10,
        leading=14,
        leftIndent=4,
    )
    
    # Standard body text — black, justified, 9pt
    normal_style = ParagraphStyle(
        'NormalCustom',
        parent=styles['Normal'],
        fontSize=9,
        leading=13,
        fontName='MicrosoftSansSerif',
        textColor=colors.HexColor("#000000"),
        alignment=TA_JUSTIFY,
        rightIndent=0,
        firstLineIndent=0
    )

    # Bold body text (RE: FEE PROPOSAL, signatory block)
    normal_style1 = ParagraphStyle(
        'NormalCustom1',
        parent=styles['Normal'],
        fontSize=9,
        leading=13,
        fontName='Roboto-Bold',
        textColor=colors.black,
        alignment=TA_JUSTIFY,
        leftIndent=0,
        rightIndent=0,
        firstLineIndent=0
    )

    # Small / footnote / note text — 8pt
    small_style = ParagraphStyle(
        'SmallStyle',
        parent=styles['Normal'],
        fontSize=8,
        leading=11,
        fontName='MicrosoftSansSerif',
        textColor=colors.HexColor("#000000"),
        bulletIndent=10,
        leftIndent=0,
        bulletFontName='Symbol',
        rightIndent=0,
        firstLineIndent=0
    )

    note_style = ParagraphStyle(
        'NoteStyle',
        parent=styles['Normal'],
        fontSize=8,
        leading=11,
        fontName='MicrosoftSansSerif',
        textColor=colors.HexColor("#000000"),
        spaceBefore=4,
        spaceAfter=4
    )

    italic_style = ParagraphStyle(
        'ItalicStyle',
        parent=styles['Normal'],
        fontSize=8,
        leading=11,
        fontName='Roboto-Italic',
        textColor=colors.HexColor("#000000"),
        spaceBefore=2,
        spaceAfter=2
    )

    italic_normal_style = ParagraphStyle(
        'ItalicNormalStyle',
        parent=styles['Normal'],
        fontSize=9,
        leading=13,
        fontName='Roboto-Italic',
        textColor=colors.HexColor("#000000"),
        spaceBefore=2,
        spaceAfter=2,
        alignment=TA_JUSTIFY
    )
    
    # ==================== PAGE 1 - COVER PAGE ====================
    cover_image_path = None
    for ext in ['cover_image.png', 'cover_image.jpg', 'cover_image.jpeg']:
        if os.path.exists(ext):
            cover_image_path = ext
            break
    
    company_name = data.get('clientCompany', 'ABC India Pvt Ltd')
    
    if cover_image_path:
        cover = CoverPageWithCompany(cover_image_path, company_name)
        elements.append(cover)
    else:
        elements.append(Spacer(1, 1.5*inch))
        elements.append(Paragraph("LEADING ASIA PACIFIC CORPORATE SOLUTIONS PROVIDER", 
                                 ParagraphStyle('cover', fontSize=10, alignment=TA_CENTER, textColor=colors.grey)))
        elements.append(Spacer(1, 0.3*inch))
        elements.append(Paragraph("INCORP GROUP PROPOSAL", ParagraphStyle('title', fontSize=22, alignment=TA_CENTER, fontName='Roboto-Bold', textColor=colors.HexColor('#C00000'))))
        elements.append(Spacer(1, 0.5*inch))
        elements.append(Paragraph(company_name, 
                                 ParagraphStyle('company', fontSize=18, alignment=TA_CENTER, 
                                              fontName='Helvetica-Bold', textColor=colors.HexColor('#002060'))))
    elements.append(PageBreak())
    
    # ==================== PAGE 5 - LETTER TO CLIENT ====================
    proposal_date = data.get('proposalDate', datetime.now().strftime('%Y-%m-%d'))
    try:
        date_obj = datetime.strptime(proposal_date, '%Y-%m-%d')
        formatted_date = date_obj.strftime('%d. %m. %Y')
    except:
        formatted_date = proposal_date

    elements.append(Paragraph(formatted_date, normal_style))
    elements.append(Spacer(1, 12))
    elements.append(Paragraph(data.get('clientName', 'Client Name'), normal_style))
    elements.append(Paragraph(data.get('clientDesignation', 'Client Designation'), normal_style))
    elements.append(Paragraph(data.get('clientCompany', 'Client Company Name'), normal_style))
    if data.get('clientAddress1'):
       elements.append(Paragraph(data.get('clientAddress1'), normal_style))
    if data.get('clientAddress2'):
       elements.append(Paragraph(data.get('clientAddress2'), normal_style))
    if data.get('clientAddress3'):
       elements.append(Paragraph(data.get('clientAddress3'), normal_style))
    if not any([data.get('clientAddress1'), data.get('clientAddress2'), data.get('clientAddress3')]):
       elements.append(Paragraph(data.get('clientAddress', 'Client Company Address'), normal_style))
    elements.append(Spacer(1, 20))
    
    elements.append(Paragraph(f"Dear {data.get('clientName', 'XXXX').split()[0]},", normal_style))
    elements.append(Spacer(1, 12))
    elements.append(Paragraph("<b><font color=\"#C00000\">RE: FEE PROPOSAL</font></b>", normal_style1))
    elements.append(Spacer(1, 8))
    
    letter_text = """We are pleased to be presenting our proposal to you.<br/><br/>
Our team of experienced professionals work very closely with clients on various corporate, accounting, compliance and governance matter and identify the unique requirements of individual organizations. As a strong believer of long-term partnerships, we are committed to providing tailored solutions that not only meet our clients' objectives, but also giving them a peace of mind to focus on their core businesses.<br/><br/>
The following pages outline our services tailor made to you and we trust that our proposal meets your expectations. We are excited to work with you and look forward to a long and mutually beneficial working relationship with you and the company.<br/><br/>
Yours Sincerely and on behalf of In.Corp,<br/><br/><br/>"""
    text="""<b>CA Bansi Shah</b><br/>
<b>Lead – International clients group</b><br/>
<b>InCorp Advisory Services Pvt Ltd</b>"""
    
    elements.append(Paragraph(letter_text, normal_style))
    elements.append(Paragraph(text, normal_style1))
    elements.append(PageBreak())
    
    # ==================== PAGE 6 - SCOPE & FEES INTRO ====================
    elements.append(Paragraph("SCOPE OF SERVICES", heading1_style))
    elements.append(Spacer(1, 1))
    
    scope_text = data.get('scopeOfServices', '')
    for line in scope_text.split('\n'):
        line = line.strip()
        if line:
            elements.append(Paragraph(f'• {line}', normal_style))
    elements.append(Spacer(1, 12))
    
    elements.append(Paragraph("FEES", heading1_style))

    fee_type = data.get('feeType', '')

    if fee_type == 'setup':
        fee_line = "Our fee structure includes initial setup fees."
    elif fee_type == 'ongoing':
        fee_line = "Our fee structure includes ongoing charges that may be billed monthly, quarterly, or annually."
    else:
        fee_line = "Our fee structure includes initial setup fees."

    fees_intro = f"""This section outlines the estimated fees for InCorp's services of your company. {fee_line} Additionally, fees may be incurred based on the time spent on specific tasks or on a per-instance basis. For any additional services not encompassed by this proposal that may incur, additional charges, we will receive your approval before any work commences. Please note that all fees mentioned are in US Dollars, exclusive of the prevailing Goods and Services Tax (GST) / Value Added Tax (VAT)."""
    elements.append(Paragraph(fees_intro, normal_style))
    elements.append(Spacer(1, 6))
    
    # ==================== DYNAMIC SECTION LETTER ASSIGNMENT ====================
    section_keys = [
    ('handover',      data.get('sectionA') == 'on' and any([data.get('includeHandover') == 'on', data.get('includeDueDiligence') == 'on'])),
    ('incorporation', data.get('sectionB') == 'on' and any([data.get('includeIncorporation') == 'on', data.get('includeGST') == 'on',
                           data.get('includeFCGPR') == 'on', data.get('includeROC') == 'on',
                           data.get('includeIEC') == 'on', data.get('includePT') == 'on',
                           data.get('includeBEN') == 'on', data.get('includeMGT') == 'on',
                           data.get('includePAN') == 'on', data.get('includeTrademark') == 'on',
                           data.get('includeForeignPAN') == 'on', data.get('includeBankAssist') == 'on',
                           data.get('includeRegOffice') == 'on', data.get('includeNomineeDir') == 'on'])),
    ('accounting',    data.get('sectionC') == 'on' and any([data.get('includeAdvanceTax') == 'on', data.get('includeTDS') == 'on',
                           data.get('includeIncomeTax') == 'on', data.get('includeGSTComp') == 'on',
                           data.get('includeCompanyLaw') == 'on', data.get('includeRBIFiling') == 'on',
                           data.get('includeMasterFiling') == 'on', data.get('includeAcctSetup') == 'on',
                           data.get('includeAcctMaint') == 'on', data.get('includeFinStmt') == 'on',
                           data.get('includePayrollSetup') == 'on', data.get('includeShopPOSH') == 'on',
                           data.get('includePayrollProc') == 'on', data.get('includeLabourLaw') == 'on',
                           data.get('includeAnnualReturns') == 'on'])),
    ('transfer',      data.get('sectionD') == 'on' and any([data.get('includeBenchmarking') == 'on', data.get('includeIntercompany') == 'on'])),
]
    letters = {}
    counter = 0
    for key, is_selected in section_keys:
        if is_selected:
            letters[key] = chr(65 + counter)
            counter += 1

    # ==================== PAGE 7 - A. HANDOVER SERVICE ====================
    if 'handover' in letters:
       if not (data.get('includeHandover') == 'on' or data.get('includeDueDiligence') == 'on'):
            pass  # skip entirely
       else:
        elements.append(Paragraph(f"{letters['handover']}. One time Handover Service", heading2_style))
        elements.append(Spacer(1, 1))
        handover_intro = f"""Since the company has been in existence since {data.get('companyYear', 'YYYY')}, we shall need to undertake a handover of the current financial, secretarial, payroll and other records of the company from current service provider."""
        elements.append(Paragraph(handover_intro, normal_style))
        elements.append(Spacer(1, 10))

        handover_data = [['Services', 'Frequency', 'Fee (In USD)']]

        if data.get('includeHandover') == 'on':
            handover_fee = format_currency(data.get('handoverFee', '0'))
            if not handover_fee:
                handover_fee = '500 '
            handover_freq = data.get('handoverFrequency', 'One-time')
            handover_data.append([
                Paragraph("""<font color="#C00000" face="Roboto-Bold"><b>Handover from erstwhile service provider of various records under laws as mentioned below.
This process does not entail conducting a due diligence.</b></font><br/>
• GST laws/regulations<br/>
• Income Tax Act, 1961<br/>
• Company's Act, 2013<br/>
• Foreign Exchange Rules & Regulations""", normal_style),
                handover_freq,
                handover_fee
            ])

        if data.get('includeDueDiligence') == 'on':
            dd_fee = format_currency(data.get('dueDiligenceFee', '0'))
            if not dd_fee:
                dd_fee = '500 per year'
            else:
                dd_fee = dd_fee + ' per year'
            dd_freq = data.get('ddFrequency', 'One-time')
            handover_data.append([
                Paragraph("""<font color="#C00000" face="Roboto-Bold">Basic due diligence from perspective of*–</font><br/>
• Company's Act, 2013<br/>
• Income Tax Act, 1961<br/>
• Goods and Service Tax Act, 2017<br/>
• Foreign Exchange Management Act, 1999""", normal_style),
                dd_freq,
                dd_fee
            ])

        if len(handover_data) > 1:
            handover_table = Table(handover_data, colWidths=[4.4*inch, 1.5*inch, 1.5*inch])
            handover_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor("#FFFFFF")),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.black),
                ('FONTNAME', (0, 0), (-1, 0), 'MicrosoftSansSerif'),
                ('FONTSIZE', (0, 0), (2, 0), 10),
                ('FONTSIZE', (0, 1), (-1, -1), 9),
                ('ALIGN', (1, 0), (2, -1), 'CENTRE'),
                ('GRID', (0, 1), (-1, -1), 0.5, colors.black),
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ('TOPPADDING', (0, 0), (-1, -1), 8),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
                ('LEFTPADDING', (0, 1), (0, -1), 2),
                ('RIGHTPADDING', (0, 0), (-1, -1), 8),
                ('LINEBELOW', (0, 0), (-1, 0), 0.75, colors.black),
                ('LINEBEFORE', (0, 0), (-1, 0), 0, colors.white),
                ('LINEAFTER', (0, 0), (-1, 0), 0, colors.white),
                ('LINEABOVE', (0, 0), (-1, 0), 0, colors.white),
            ]))
            elements.append(handover_table)
            elements.append(Spacer(1, 10))
            elements.append(Paragraph("*Any fees for rectification (or) completion of pending past compliances shall attract additional fees and we shall seek your approval prior to commencement of that work.", italic_style))
            elements.append(Spacer(1, 10))

        notes_a = """<b><u><font color="#002060">Note:</font></u></b><br/><br/>
• All fees quoted above exclude 18% GST<br/>
• Professional fees exclude any fees towards regularisation of past non compliances.<br/>
• Advance of 100% of the above selected option.<br/><br/>"""
        elements.append(Paragraph(notes_a, normal_style))
        elements.append(Paragraph("* Any other services not specifically quoted above and not specifically agreed separately shall be chargeable as under:", italic_style))
        elements.append(Spacer(1, 4))
        elements.append(Paragraph("<b>For Partner: USD 300 per Hour</b>", italic_style))
        elements.append(Spacer(1, 2))
        elements.append(Paragraph("<b>For Associates: USD 200 per Hour</b>", italic_style))
    
    
    # ==================== PAGE 8 - B. INCORPORATION SERVICE ====================
    if 'incorporation' in letters:
        elements.append(Spacer(1, 12))
        elements.append(Paragraph(f"{letters['incorporation']}. Incorporation / Secretarial Service and Mandatory Registrations post Incorporation", heading2_style))
        elements.append(Spacer(1, 2))

        inc_data = [['Services', 'One-time Fee']]

        if data.get('includeIncorporation') == 'on':
            inc_fee = format_currency(data.get('incorporationFee', '0'))
            if not inc_fee:
                inc_fee = '1500'
            inc_data.append([
                Paragraph("""<font color="#C00000" face="Roboto-Bold"><b>Incorporation</b></font><br/>
• PAN of the company included<br/>
• TAN of the company included<br/>
• Employees' Provident Fund and Miscellaneous Provision Act, Employees' State Insurance Corporation Act included""", normal_style),
                inc_fee
            ])

        if data.get('includeGST') == 'on':
            gst_fee = format_currency(data.get('gstRegFee', '0'))
            if not gst_fee:
                gst_fee = '350'
            inc_data.append([
                Paragraph("""<font color="#C00000" face="Roboto-Bold"><b>Goods & Service Tax (GST)</b></font><br/><br/>
Registration of single location with GST authorities.<br/><br/>
<i>Registration of every additional location with the GST authorities shall cost USD 100</i>""", normal_style),
                gst_fee
            ])

        if data.get('includeFCGPR') == 'on':
            fcgpr_fee = format_currency(data.get('fcgprFee', '0'))
            if not fcgpr_fee:
                fcgpr_fee = '1250 per applicant'
            else:
                fcgpr_fee = fcgpr_fee + ' per applicant'
            inc_data.append([
                Paragraph("""<font color="#C00000" face="Roboto-Bold"><b>FCGPR Filing with Reserve Bank of India</b></font><br/>
Filing of Forms and declaration with RBI as required under FEMA""", normal_style),
                fcgpr_fee
            ])

        if data.get('includeROC') == 'on':
            roc_fee = format_currency(data.get('rocComplianceFee', '0'))
            if not roc_fee:
                roc_fee = '500'
            inc_data.append([
                Paragraph("""<font color="#C00000" face="Roboto-Bold"><b>Statutory Compliances with Registrar of Companies under Companies Act:</b></font><br/><br/>
• Drafting of first board meeting documents<br/><br/>
• Guidance on capital infusion in bank account<br/><br/>
• File form with Ministry for commencement of business (COC)<br/><br/>
• Preparation of statutory shareholders register""", normal_style),
                roc_fee
            ])

        if len(inc_data) > 1:
            inc_table = Table(inc_data, colWidths=[5.7*inch, 1.5*inch])
            inc_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor("#FFFFFF")),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.black),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (1, 0), 10),
                ('FONTSIZE', (0, 1), (-1, -1), 8),
                ('ALIGN', (1, 0), (1, -1), 'CENTER'),
                ('GRID', (0, 1), (-1, -1), 0.5, colors.black),
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ('TOPPADDING', (0, 0), (-1, -1), 8),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
                ('LEFTPADDING', (0, 1), (0, -1), 2),
                ('RIGHTPADDING', (0, 0), (-1, -1), 8),
                ('LINEBELOW', (0, 0), (-1, 0), 0.75, colors.black),
                ('LINEBEFORE', (0, 0), (-1, 0), 0, colors.white),
                ('LINEAFTER', (0, 0), (-1, 0), 0, colors.white),
                ('LINEABOVE', (0, 0), (-1, 0), 0, colors.white),
            ]))
            elements.append(inc_table)
            elements.append(Spacer(1, 13))

        notes_b = """<b><u><font color="#002060">Note:</font></u></b><br/><br/>
• All fees quoted above exclude 18% GST.<br/>
• Professional fees exclude all out-of-pocket expenses like filing fees, courier expenses, apostilling & notary cost to any authorities/departments, statutory fees payable to Registrar of companies (ROC) towards incorporation etc. other than those mentioned above.<br/>
• Advance of 100% of the above selected option.<br/>
• On finalization of shareholding structure, we shall be able to guide on compliances needed for issuance of share certificates and shall share a separate fee quote for the same.<br/>
"""
        text_b = """<br/><b><i>* Any other services not specifically quoted above and not specifically agreed separately shall be chargeable as under</i></b><br/><br/>
<b><i>For Partner: USD 300 per Hour</i></b><br/><br/>
<b><i>For Associates: USD 200 per Hour</i></b>"""
        elements.append(Paragraph(notes_b, normal_style))
        elements.append(Paragraph("* Any other services not specifically quoted above and not specifically agreed separately shall be chargeable as under:", italic_style))
        elements.append(Spacer(1, 4))
        elements.append(Paragraph("<b>For Partner: USD 300 per Hour</b>", italic_style))
        elements.append(Spacer(1, 2))
        elements.append(Paragraph("<b>For Associates: USD 200 per Hour</b>", italic_style))
        elements.append(Spacer(1, 8))

        # ==================== PAGE 9 - OPTIONAL REGISTRATIONS ====================
        elements.append(Paragraph("Optional registrations required post incorporation (One-time)", heading2_style))
        elements.append(Spacer(1, 3))

        opt_data = [['Services', 'Fees (In USD)']]
        optional_services = [
            ('includeIEC', 'iecFee','200', '<font color="#C00000"face="Roboto-Bold">Import Export Code (IEC Code)</font>'),
            ('includePT', 'ptFee', '200','<font color="#C00000"face="Roboto-Bold">Profession Tax (PT)</font><br/><br/>•Payments and return filing for company, its employees until the company\'s certificate of commencement is obtained'),
            ('includeBEN', 'benFee','250', '<font color="#C00000" face="Roboto-Bold">Submission of for Significant Beneficial Ownership via form BEN-2</font>'),
            ('includeMGT', 'mgtFee', '250','<font color="#C00000" face="Roboto-Bold">Filing of requisite forms with Registrar of Companies (ROC) with respect to beneficial and nominee shareholding (via Form MGT 4, MGT 5, MGT 6)</font>'),
            ('includePAN', 'panCardFee','300', '<font color="#C00000" face="Roboto-Bold">Physical PAN Card of the company</font>'),
            ('includeTrademark', 'trademarkFee','350', '<font color="#C00000" face="Roboto-Bold">Trademark Registration (exclusive of disbursement fees)</font>'),
            ('includeForeignPAN', 'foreignPanFee','200 per director', '<font color="#C00000" face="Roboto-Bold">PAN for foreign director</font>'),
            ('includeBankAssist', 'bankAssistFee','250', '<font color="#C00000" face="Roboto-Bold">Assistance in opening of bank account</font>')
        ]
        opt_style = ParagraphStyle('OptStyle', parent=normal_style, leftIndent=0, rightIndent=0)
        per_director_keys = {'foreignPanFee'}
        for checkbox, fee_key, default_fee, label in optional_services:
            if data.get(checkbox) == 'on':
                fee = format_currency(data.get(fee_key, '0'))
                if not fee:
                    fee = default_fee
                else:
                    # Append suffix if this fee type needs it
                    if fee_key in per_director_keys:
                        fee = fee + ' per director'
                opt_data.append([Paragraph(label, opt_style), fee])

        if len(opt_data) > 1:
            opt_table = Table(opt_data, colWidths=[5.7*inch, 1.5*inch])
            opt_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor("#FFFFFF")),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.black),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (1, 0), 10),
                ('FONTSIZE', (0, 1), (-1, -1), 8),
                ('ALIGN', (1, 0), (1, -1), 'CENTER'),
                ('GRID', (0, 1), (-1, -1), 0.5, colors.black),
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ('TOPPADDING', (0, 0), (-1, -1), 8),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
                ('LEFTPADDING', (0, 1), (0, -1), 2),
                ('RIGHTPADDING', (0, 0), (-1, -1), 8),
                ('LINEBELOW', (0, 0), (-1, 0), 0.75, colors.black),
                ('LINEBEFORE', (0, 0), (-1, 0), 0, colors.white),
                ('LINEAFTER', (0, 0), (-1, 0), 0, colors.white),
                ('LINEABOVE', (0, 0), (-1, 0), 0, colors.white),
            ]))
            elements.append(opt_table)
            elements.append(Spacer(1, 10))

        elements.append(Paragraph("*For every new director's professional tax no., there shall be additional cost of $100 per director", italic_style))
        elements.append(Paragraph("*Digital signature certificate (DSC) token can be obtained at a cost of USD 200 per applicant.", italic_style))
        elements.append(Spacer(1, 20))

        notes_opt = """<b><u><font color="#002060">Note:</font></u></b><br/><br/>
• All fees quoted above exclude 18% GST.<br/>
• Professional fees exclude all out-of-pocket expenses like filing fees, courier expenses, apostilling &amp; notary cost to any authorities/departments, statutory fees payable to Registrar of companies (ROC) towards incorporation etc. other than those mentioned above.<br/>
• Advance of 100% of the above selected option.<br/>
"""
        elements.append(Paragraph(notes_opt, normal_style))
        elements.append(Paragraph("*Any other services not specifically quoted above and not specifically agreed separately shall be chargeable as under:", italic_style))
        elements.append(Spacer(1, 4))
        elements.append(Paragraph("<b>For Partner: USD 300 per Hour</b>", italic_style))
        elements.append(Spacer(1, 2))
        elements.append(Paragraph("<b>For Associates: USD 200 per Hour</b>", italic_style))
        elements.append(Spacer(1, 4))

        # NOMINEE DIRECTOR SERVICE
        elements.append(Paragraph("Nominee Director and Registered Office Address Service", heading2_style))
        elements.append(Spacer(1, 4))

        nominee_data = [['Services', 'Monthly Fee(in USD)']]

        if data.get('includeRegOffice') == 'on':
            reg_office_fee = format_currency(data.get('registeredOfficeFee', '0'))
            if not reg_office_fee:
                reg_office_fee = '300'
            nominee_data.append([
                Paragraph("""<font color="#C00000" face="Roboto-Bold"><b>Registered Office Service</b></font><br/><br/>
A refundable Security deposit @USD 2500 applies**. Refundable upon cessation of Registered office service.""", normal_style),
                reg_office_fee
            ])

        if data.get('includeNomineeDir') == 'on':
            nom_dir_fee = format_currency(data.get('nomineeDirectorFee', '0'))
            if not nom_dir_fee:
                nom_dir_fee = '350'
            nominee_data.append([
                Paragraph("""<font color="#C00000" face="Roboto-Bold"><b>Nominee Director Service</b></font><br/>
A refundable Security deposit per nominee @USD 5000 applies*. Refundable upon cessation of Nominee Director Service<br/><br/>
Director's fee for attending a physical or recorded or live board meeting @USD300 per director per board meeting<br/><br/>
Every nominee director needs to be protected under a director's indemnity policy. Premium of indemnity bond to be charged on actual basis. InCorp shall enter into a separate nominee directors' agreement at the time of engagement.<br/><br/>
To ensure the removal of a nominee director from registrations ***with various authorities where required, InCorp must be notified at least three months in advance. Additionally, professional fees for this service will continue to be charged until the removal is reflected by all relevant authorities as well as Bank & new director is appointed in his place.
""", normal_style),
                nom_dir_fee
            ])

        if len(nominee_data) > 1:
            nominee_table = Table(nominee_data, colWidths=[5.7*inch, 1.5*inch])
            nominee_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor("#FFFCFC")),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.black),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (1, 0), 10),
                ('FONTSIZE', (0, 1), (-1, -1), 8),
                ('ALIGN', (1, 0), (1, -1), 'CENTER'),
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ('GRID', (0, 1), (-1, -1), 0.5, colors.black),
                ('TOPPADDING', (0, 0), (-1, -1), 5),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 5),
                ('LEFTPADDING', (0, 1), (0, -1), 2),
                ('RIGHTPADDING', (0, 0), (-1, -1), 8),
                ('LINEBELOW', (0, 0), (-1, 0), 0.75, colors.black),
                ('LINEBEFORE', (0, 0), (-1, 0), 0, colors.white),
                ('LINEAFTER', (0, 0), (-1, 0), 0, colors.white),
                ('LINEABOVE', (0, 0), (-1, 0), 0, colors.white),
            ]))
            elements.append(nominee_table)
            elements.append(Paragraph("*Failure to engage InCorp's services for regular compliances of the company post the setup such as tax, secretarial, FEMA etc. shall result in forfeiture of the security deposit received against nominee director and registered office services.", italic_style))
            elements.append(Spacer(1, 4))
            elements.append(Paragraph("**Any fees for rectification (or) completion of pending past compliances shall attract additional fees and we shall seek your approval prior to commencement of that work.", italic_style))
            elements.append(Spacer(1, 4))
            elements.append(Paragraph("*** The Nominee Director shall not sign any return, forms or documents relating to any statutory filing nor will be appointed as the authorized signatory to any of the bank accounts of the entity or under GST, Income Tax any other government portal. The Company may consider appointing one of its key managerial personnel as the authorised signatory across all government portals", italic_style))

        # ==================== PAGE 10 - NOMINEE NOTES ====================
            elements.append(Spacer(1, 10))
            nominee_notes = """<b><u><font color="#002060">Note:</font></u></b><br/><br/>
• All fees quoted above exclude 18% GST.<br/>
• The Nominee Director will not be involved in day-to-day affairs / management of the Company. He/She shall not sign any return, forms or documents relating to any statutory filing.<br/>
• The service of Registered office & Nominee director is offered on discretionary basis only for temporary basis of 6 months. Such services are provided only in case of successful completion of Internal Customer Due diligence at InCorp and formal engagement of Incorp for all the compliances (tax, secretarial, FEMA etc.) post incorporation of the company for the regular maintenance of the company.<br/>
• Failure to engage InCorp's services for regular compliances of the company post the setup such as tax, secretarial, FEMA etc. shall result in forfeiture of the security deposit received against registered office and nominee director services.<br/>
• Professional fees exclude all out-of-pocket expenses like filing fees, courier expenses, apostilling & notary cost to any authorities/departments, statutory fees payable to Registrar of companies (ROC) towards incorporation etc. other than those mentioned above.<br/>
• Advance of 100% of the above selected option.<br/>
"""
            elements.append(Paragraph(nominee_notes, normal_style))
            elements.append(Paragraph("* Any other services not specifically quoted above shall be chargeable as under:", italic_style))
            elements.append(Spacer(1, 4))
            elements.append(Paragraph("<b>For Partner: USD 300 per Hour</b>", italic_style))
            elements.append(Spacer(1, 2))
            elements.append(Paragraph("<b>For Associates: USD 200 per Hour</b>", italic_style))

    # ==================== PAGE 11-12 - C. ACCOUNTING / TAX / PAYROLL ====================
    if 'accounting' in letters:
        elements.append(Spacer(1, 12))
        elements.append(Paragraph(f"{letters['accounting']}. Accounting / Tax / Payroll / Annual Compliance Services", heading2_style))
        elements.append(Spacer(1, 5))
        acc_intro = """The below quotation is our base fees for first year of business with limited volume of transactions and may change depending upon volume of work and nature of transactions:"""
        elements.append(Paragraph(acc_intro, small_style))
        elements.append(Spacer(1, 12))

        # ONE BIG COMBINED TABLE
        all_sections_data = [['Services', 'Frequency', 'Notes', 'Fees (in USD)']]

        # Track for totals calculation
        total_annual_cost = 0
        total_onetime_cost = 0

        def add_to_totals(frequency, fee_str):
            """Helper function to add fees to totals based on frequency"""
            nonlocal total_annual_cost, total_onetime_cost
            try:
                fee = float(fee_str.replace(',', '')) if fee_str else 0
                freq_lower = frequency.lower()
                if 'one time' in freq_lower or 'one-time' in freq_lower:
                    total_onetime_cost += fee
                elif 'monthly' in freq_lower:
                    total_annual_cost += fee * 12
                elif 'quarterly' in freq_lower:
                    total_annual_cost += fee * 4
                elif 'annual' in freq_lower:
                    total_annual_cost += fee
            except:
                pass

    # Initialize accounting variables always (needed even if accounting section not selected)
    if 'accounting' not in letters:
        all_sections_data = [['Services', 'Frequency', 'Notes', 'Fees (in USD)']]
        total_annual_cost = 0
        total_onetime_cost = 0
        def add_to_totals(frequency, fee_str):
            pass
    
    section_header_style = ParagraphStyle(
        'SectionHeader',
        parent=normal_style,
        fontSize=9,
        leading=13,
        fontName='Roboto-Bold',
        textColor=colors.HexColor('#C00000'),
        leftIndent=0,
        rightIndent=0,
        wordWrap='LTR'
    )
    # DIRECT TAX
    direct_tax_entries = []
    if data.get('includeAdvanceTax') == 'on':
        advance_tax_fee = format_currency(data.get('advanceTaxFee', '0'))
        if not advance_tax_fee:
            advance_tax_fee = '200 per month'
        freq = data.get('advanceTaxFrequency', 'Quarterly')
        add_to_totals(freq, advance_tax_fee)
        direct_tax_entries.append([
            '',
            freq,
            Paragraph("1) Advance tax Compliances<br/> • Quarterly calculations and payment", normal_style),
            advance_tax_fee
        ])
    
    if data.get('includeTDS') == 'on':
        tds_fee = format_currency(data.get('tdsFee', '0'))
        if not tds_fee:
            tds_fee = '200 per month'
        freq = data.get('tdsFrequency', 'Monthly/Quarterly')
        add_to_totals(freq, tds_fee)
        direct_tax_entries.append([
            '',
            freq,
            Paragraph("""2) TDS compliances:<br/>
• Calculation and Payment of TDS<br/>
• Filing of TDS Returns<br/><br/>
(The above excludes cost of revisions of TDS returns)""", normal_style),
            tds_fee
        ])
    
    if data.get('includeIncomeTax') == 'on':
        income_tax_fee = format_currency(data.get('incomeTaxReturnFee', '0'))
        if not income_tax_fee:
            income_tax_fee = '500 per annum'
        freq = data.get('incomeTaxFrequency', 'Annual')
        add_to_totals(freq, income_tax_fee)
        direct_tax_entries.append([
            '',
            freq,
            Paragraph("""3) Annual Income tax return<br/>
Computation and filing of Annual Income tax Return<br/><br/>
4) Statement of Financial Transactions (SFT) – Basic Reporting""", normal_style),
            income_tax_fee
        ])
       

    
    if direct_tax_entries:
        all_sections_data.append([
            Paragraph('<font color="#C00000" face="Roboto-Bold"><b>Direct tax compliances</b></font>',  section_header_style),
            direct_tax_entries[0][1],
            direct_tax_entries[0][2],
            direct_tax_entries[0][3]
        ])
        for entry in direct_tax_entries[1:]:
            all_sections_data.append(entry)
    
    # INDIRECT TAX
    if data.get('includeGSTComp') == 'on':
        gst_fee = format_currency(data.get('gstComplianceFee', '0'))
        if not gst_fee:
            gst_fee = '250 per month'
        freq = data.get('gstFrequency', 'Monthly and Annual')
        add_to_totals(freq, gst_fee)
        all_sections_data.append([
            Paragraph('<font color="#C00000" face="Roboto-Bold"><b>Indirect tax compliances</b></font>', section_header_style),
            freq,
            Paragraph("""1) GST Compliances:<br/>
• Calculations and payment of GST<br/>
• Filing of monthly GST Returns<br/><br/>
(The above excludes cost of revisions of TDS returns. The fee quote for the same shall be shared if applicable) note should be in Italics)""", normal_style),
            gst_fee
        ])
    
    # COMPANY LAW
    if data.get('includeCompanyLaw') == 'on':
        company_law_fee = format_currency(data.get('companyLawFee', '0'))
        if not company_law_fee:
            company_law_fee = '200 per month'
        freq = data.get('companyLawFrequency', 'Monthly')
        add_to_totals(freq, company_law_fee)
        all_sections_data.append([
            Paragraph('<font color="#C00000" face="Roboto-Bold"><b>Company Law</b></font>', normal_style),
            freq,
            Paragraph("""Company Law Compliances (Scope as per Annexure 1)<br/>
Assistance on conduction of virtual board meeting – USD 150 per board meeting""",normal_style),
            company_law_fee
        ])
    
    # FOREIGN EXCHANGE
    foreign_exchange_entries = []
    if data.get('includeRBIFiling') == 'on':
        rbi_filing_fee = format_currency(data.get('rbiFilingFee', '0'))
        if not rbi_filing_fee:
            rbi_filing_fee = '450 per annum'
        freq = data.get('rbiFilingFrequency', 'Annual')
        add_to_totals(freq, rbi_filing_fee)
        foreign_exchange_entries.append([
            '',
            freq,
            Paragraph("Annual Filings with Reserve bank of India", normal_style),
            rbi_filing_fee
        ])
    
    if data.get('includeMasterFiling') == 'on':
        master_filing_fee = format_currency(data.get('masterFilingFee', '0'))
        if not master_filing_fee:
            master_filing_fee = '350 per annum'
        freq = data.get('masterFilingFrequency', 'Annual')
        add_to_totals(freq, master_filing_fee)
        foreign_exchange_entries.append([
            '',
            freq,
            Paragraph("Annual Master Filing Form 3CEAA Part A (Basic Reporting)", normal_style),
            master_filing_fee
        ])
    
    if foreign_exchange_entries:
        all_sections_data.append([
            Paragraph('<font color="#C00000" face="Roboto-Bold"><b>Foreign Exchange laws</b></font>', normal_style),
            foreign_exchange_entries[0][1],
            foreign_exchange_entries[0][2],
            foreign_exchange_entries[0][3]
        ])
        for entry in foreign_exchange_entries[1:]:
            all_sections_data.append(entry)
    
    # ACCOUNTING (continues on same page - no PageBreak)
    accounting_entries = []
    if data.get('includeAcctSetup') == 'on':
        acct_setup_fee = format_currency(data.get('accountingSetupFee', '0'))
        if not acct_setup_fee:
            acct_setup_fee = '300 one time'
        freq = data.get('acctSetupFrequency', 'One time')
        add_to_totals(freq, acct_setup_fee)
        accounting_entries.append([
            '',
            freq,
            Paragraph("""Setup of accounting software<br/>
• Liaison with the software expert for the setup<br/>
• Ensure due configuration of the software with applicable laws<br/>
• Short tutorial on guidance with respect to use of accounting software""", normal_style),
            acct_setup_fee
        ])
    
    if data.get('includeAcctMaint') == 'on':
        accounting_tiers = data.get('accountingEntries', [])
        
        # Default values if not provided
        if not accounting_tiers or len(accounting_tiers) == 0:
            accounting_tiers = [
                {'transactions': 'Upto 20', 'fee': '200'},
                {'transactions': '20 - 50', 'fee': '250'},
                {'transactions': '50-80', 'fee': '300'}
            ]
        
        nested_table_data = [
            [Paragraph('<b>No. of transactions per month</b>', ParagraphStyle('nested', fontSize=8, alignment=TA_LEFT)), 
             Paragraph('<b>Fees per month (in USD)</b>', ParagraphStyle('nested', fontSize=8, alignment=TA_RIGHT))]
        ]
        
        for tier in accounting_tiers:
            transactions = tier.get('transactions', '0')
            fee = tier.get('fee', '0')
            nested_table_data.append([transactions, fee])
        
        nested_table = Table(nested_table_data, colWidths=[1.4*inch, 1.4*inch])
        nested_table.setStyle(TableStyle([
            ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
            ('FONTSIZE', (0, 0), (-1, -1), 8),
            ('ALIGN', (1, 0), (1, -1), 'RIGHT'),        
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('TOPPADDING', (0, 0), (-1, -1), 3),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 3),
            ('LEFTPADDING', (0, 0), (-1, -1), 3),
            ('RIGHTPADDING', (0, 0), (-1, -1), 3),
            ('NOSPLIT', (0, 0), (-1, -1)),
        ]))
        
        notes_para = Paragraph("""Accounting and maintenance of books of accounts:<br/>
• Data entry in accounting software<br/>
• Weekly processing of Bank Reconciliation<br/>
• Weekly processing of Purchase invoices<br/>
• Maker access in bank account/preparing payments<br/>
• Weekly forwarding of open suppliers/customers<br/>
• Preparation of Monthly Profit & loss Statement and Balance Sheet""", normal_style)
        
        freq = data.get('acctMaintFrequency', 'Monthly')
        
        # Get the monthly fee from form (this is the fee that shows in last column)
        accounting_monthly_fee = format_currency(data.get('accountingMaintenanceFee', '200'))
        if not accounting_monthly_fee:
            accounting_monthly_fee = '200'
        
        # Add monthly fee to totals calculation
        add_to_totals(freq, accounting_monthly_fee)
        
        accounting_entries.append([
            '',
            freq,
            [notes_para, Spacer(1, 4), nested_table],
            f'{accounting_monthly_fee} per month'
        ])
    
    if data.get('includeFinStmt') == 'on':
        fin_stmt_fee = format_currency(data.get('financialStatementsFee', '0'))
        if not fin_stmt_fee:
            fin_stmt_fee = '500 per annum'
        freq = data.get('finStmtFrequency', 'Annual')
        add_to_totals(freq, fin_stmt_fee)
        accounting_entries.append([
            '',
            freq,
            Paragraph("""• Preparation of the financial Statements as per the Indian accounting Standards<br/>
• Liaising with auditors for audit, compliance and related matters""", normal_style),
            fin_stmt_fee
        ])
    
    if accounting_entries:
        all_sections_data.append([
            Paragraph('<font color="#C00000" face="Roboto-Bold"><b>Accounting</b></font>', normal_style),
            accounting_entries[0][1],
            accounting_entries[0][2],
            accounting_entries[0][3]
        ])
        for entry in accounting_entries[1:]:
            all_sections_data.append(entry)
    
    # PAYROLL
    payroll_entries = []
    if data.get('includePayrollSetup') == 'on':
        payroll_setup_fee = format_currency(data.get('payrollSetupFee', '0'))
        if not payroll_setup_fee:
            payroll_setup_fee = '500 one time'
        freq = data.get('payrollSetupFrequency', 'One time')
        add_to_totals(freq, payroll_setup_fee)
        payroll_entries.append([
            '',
            freq,
            'Payroll Setup (Scope as per Annexure 2)',
            payroll_setup_fee
        ])
    
    if data.get('includeShopPOSH') == 'on':
        shop_posh_fee = format_currency(data.get('shopPOSHFee', '0'))
        if not shop_posh_fee:
            shop_posh_fee = '0'
        freq = data.get('shopPOSHFrequency', 'One time')
        add_to_totals(freq, shop_posh_fee)
        payroll_entries.append([
            '',
            freq,
            Paragraph("""1. Obtaining Shop and establishment registration under Karnataka Shop and establishment act<br/>
2. Drafting of POSH (Prevention of Sexual Harassment at Workplace) policy""", normal_style),
            shop_posh_fee
        ])
    
    if data.get('includePayrollProc') == 'on':
        payroll_tiers = data.get('payrollEntries', [])
        
        # Default values if not provided
        if not payroll_tiers or len(payroll_tiers) == 0:
            payroll_tiers = [
                {'employees': 'Upto 10 employees', 'amount': '125 USD'},
                {'employees': '11 - 20 employees', 'amount': '200 USD'}
            ]
        
        nested_payroll_data = [
            [Paragraph('<b>No of employees</b>', ParagraphStyle('nested', fontSize=8, alignment=TA_LEFT)), 
             Paragraph('<b>Amount in USD per month</b>', ParagraphStyle('nested', fontSize=8, alignment=TA_RIGHT))]
        ]
        
        for tier in payroll_tiers:
            employees = tier.get('employees', '')
            amount = tier.get('amount', '')
            nested_payroll_data.append([employees, amount])
        
        nested_payroll_table = Table(nested_payroll_data, colWidths=[1.4*inch, 1.4*inch])
        nested_payroll_table.setStyle(TableStyle([
            ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
            ('FONTSIZE', (0, 0), (-1, -1), 8),
            ('ALIGN', (1, 0), (1, -1), 'RIGHT'),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('TOPPADDING', (0, 0), (-1, -1), 3),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 3),
            ('LEFTPADDING', (0, 0), (-1, -1), 3),
            ('RIGHTPADDING', (0, 0), (-1, -1), 3),
        ]))
        
        payroll_notes_para = Paragraph("Payroll Processing** (Scope as per Annexure 3)", normal_style)
        
        freq = data.get('payrollProcFrequency', 'Monthly')
        
        # Get the monthly fee from form (this is the fee that shows in last column)
        payroll_monthly_fee = format_currency(data.get('payrollProcessingFee', '125'))
        if not payroll_monthly_fee:
            payroll_monthly_fee = '125'
        
        # Add monthly fee to totals calculation
        add_to_totals(freq, payroll_monthly_fee)
        
        payroll_entries.append([
            '',
            freq,
            [payroll_notes_para, Spacer(1, 4), nested_payroll_table],
            f'{payroll_monthly_fee} per month'
        ])
    
    if data.get('includeLabourLaw') == 'on':
        labour_fee = format_currency(data.get('labourLawFee', '0'))
        if not labour_fee:
            labour_fee = '200 per month'
        freq = data.get('labourLawFrequency', 'Monthly')
        add_to_totals(freq, labour_fee)
        payroll_entries.append([
            '',
            freq,
            Paragraph("""Labour Law Compliances • Payments and return filing under:<br/>
• Provident Fund<br/>
• Employees State Insurance Corporation<br/>
• Profession Tax<br/>
•Labor Welfare Fund<br/>
(for employees upto 20 – fixed fee)""", normal_style),
            labour_fee
        ])
    
    if data.get('includeAnnualReturns') == 'on':
        annual_ret_fee = format_currency(data.get('annualReturnsFee', '0'))
        if not annual_ret_fee:
            annual_ret_fee = '200 per annum'
        freq = data.get('annualReturnsFrequency', 'Annual')
        add_to_totals(freq, annual_ret_fee)
        payroll_entries.append([
            '',
            freq,
            Paragraph("""Annual Return under the following labor law compliances:<br/>
• Sexual Harassment of Women at Workplace Act, 2013<br/>
• Shop and Establishment Act<br/>
• Maternity Act<br/>
• Gratuity Act""", normal_style),
            annual_ret_fee
        ])
    
    if payroll_entries:
        all_sections_data.append([
            Paragraph('<font color="#C00000" face="Roboto-Bold"><b>Payroll</b></font>', normal_style),
            payroll_entries[0][1],
            payroll_entries[0][2],
            payroll_entries[0][3]
        ])
        for entry in payroll_entries[1:]:
            all_sections_data.append(entry)
    total_label_style = ParagraphStyle(
    'TotalLabelStyle',
    parent=small_style,
    alignment=TA_CENTER,
)
    # ADD TOTAL ROWS
    all_sections_data.append([
        Paragraph('<b>Total costs (excluding one time costs)</b>', total_label_style),
        '',
        '',
        f'{int(total_annual_cost):,} per annum'
    ])
    
    all_sections_data.append([
        Paragraph('<b>One-time costs</b>', total_label_style),
        '',
        '',
        f'{int(total_onetime_cost):,} one time'
    ])
    
    # CREATE ONE BIG TABLE WITH ALL SECTIONS
    if len(all_sections_data) > 1:
        all_sections_table = Table(all_sections_data, colWidths=[1.3*inch, 1.3*inch, 3.2*inch, 1.4*inch])
        
        
        table_style_all = [
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor("#FFFFFF")),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.black),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
             ('FONTSIZE', (0, 0), (3, 0), 10),
            ('ALIGN', (1, 0), (3, -1), 'CENTER'),
            ('GRID', (0, 1  ), (-1, -1), 0.5, colors.black),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ('TOPPADDING', (0, 0), (-1, -1), 4),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
            ('FONTSIZE', (1, 1), (1, -1), 9),
            ('LINEBELOW', (0, 0), (-1, 0), 0.75, colors.black),  # ✅ sirf bottom border
('LINEBEFORE', (0, 0), (-1, 0), 0, colors.white),    # ✅ left border NONE
('LINEAFTER', (0, 0), (-1, 0), 0, colors.white),     # ✅ right border NONE
('LINEABOVE', (0, 0), (-1, 0), 0, colors.white), 
('FONTNAME', (1, 1), (1, -1), 'Helvetica'),

# ✅ Fees column (200 per month, 0) - BOLD + BIGGER  
('FONTSIZE', (3, 1), (3, -1), 9),
('FONTNAME', (3, 1), (3, -1), 'Helvetica'),
        ]
        
        # Add spans
        current_row = 1
        
        if len(direct_tax_entries) > 1:
            table_style_all.append(('SPAN', (0, current_row), (0, current_row + len(direct_tax_entries) - 1)))
        current_row += len(direct_tax_entries)
        
        if data.get('includeGSTComp') == 'on':
            current_row += 1
        
        if data.get('includeCompanyLaw') == 'on':
            current_row += 1
        
        if len(foreign_exchange_entries) > 1:
            table_style_all.append(('SPAN', (0, current_row), (0, current_row + len(foreign_exchange_entries) - 1)))
        current_row += len(foreign_exchange_entries)
        
        if len(accounting_entries) > 1:
            table_style_all.append(('SPAN', (0, current_row), (0, current_row + len(accounting_entries) - 1)))
        current_row += len(accounting_entries)
        
        if len(payroll_entries) > 1:
            table_style_all.append(('SPAN', (0, current_row), (0, current_row + len(payroll_entries) - 1)))
        current_row += len(payroll_entries)
        
        # Style for total rows (last 2 rows)
        total_rows_start = len(all_sections_data) - 2
        table_style_all.append(('SPAN', (0, total_rows_start), (2, total_rows_start)))
        table_style_all.append(('SPAN', (0, total_rows_start + 1), (2, total_rows_start + 1)))
        table_style_all.append(('BACKGROUND', (0, total_rows_start), (-1, -1), colors.HexColor("#FFFFFF")))
        
        all_sections_table.setStyle(TableStyle(table_style_all))
        elements.append(all_sections_table)
    
    if 'accounting' in letters:
        elements.append(Spacer(1, 10))
        elements.append(Paragraph("*The above quotation fee is for approx.20 transactions per month", italic_style))

        # ==================== PAGE 13 - NOTES ====================
        elements.append(Spacer(1, 4))
        elements.append(Paragraph("""^InCorp's empanelled audit partners can offer the services of statutory audit (applicable to all), tax audit 
(applicable on if Turnover exceeds Rs. 100 Mn) and GST audit services (If Turnover exceeds Rs. 50 Mn) and 
transfer pricing reporting & audit (applicable for companies having intercompany transactions). The quotes for 
the same can be provided separately.""", italic_style))
        elements.append(Spacer(1, 6))
        elements.append(Paragraph("""^Audit partner firms (Jayesh Sanghrajka &amp; Associates, Manish Modi &amp; Associates) shall be able to assist on that 
front. The estimated statutory fee quote for the first FY shall be between USD 2500 TO USD 3500. The auditor 
shall be able to provide the final fee quote closer to year end March 2026 depending on the nature and 
complexity of transactions.""", italic_style))
        elements.append(Spacer(1, 4))
        notes_c = """<b><u><font color="#002060">Note:</font></u></b><br/><br/>
• All fees quoted above exclude 18% GST.<br/>
• Professional fees exclude all out-of-pocket expenses like filing fees, courier expenses, government/statutory fees etc.<br/>
• Advance of 100% of the above selected option<br/>"""
        elements.append(Paragraph(notes_c, normal_style))
        elements.append(Spacer(1, 6))
        elements.append(Paragraph("*** Any other services not specifically quoted above and not specifically agreed separately shall be chargeable as under:", italic_style))
        elements.append(Spacer(1, 4))
        elements.append(Paragraph("<b>For Partner: USD 300 per Hour</b>", italic_style))
        elements.append(Spacer(1, 2))
        elements.append(Paragraph("<b>For Associates: USD 200 per Hour</b>", italic_style))
        elements.append(Spacer(1, 15))

    # ==================== TRANSFER PRICING ====================
    if 'transfer' in letters:
        elements.append(Paragraph(f"{letters['transfer']}. Transfer Pricing compliances", heading2_style))
        elements.append(Spacer(1, 4))

        tp_data = [['Services', 'Frequency', 'Notes', 'Fee (In USD)']]
        benchmark_fee = None
        interco_fee = None

        if data.get('includeBenchmarking') == 'on':
            benchmark_fee = format_currency(data.get('benchmarkingFee', '0'))
            if not benchmark_fee:
                benchmark_fee = '1500 per business activity'
            else:
                benchmark_fee = benchmark_fee + ' per business activity'
            tp_data.append([
                Paragraph('<font color="#C00000" face="Roboto-Bold"><b>Benchmarking</b></font>', normal_style),
                'One-time',
                Paragraph("""1. Assistance in conducting Functional, Asset and Risk Analysis of the proposed transaction to be entered between related parties.<br/>
2. Assisting in arriving at the arm's length price or margin range that may be applicable to the proposed transaction. Arm's Length is price that Indian &lt;company name&gt; would have charged any other non related party/clients globally for similar services. This is a legal requirement from Indian Income tax to ensure Indian revenue department is not a loss of tax revenue. and<br/>
Preparation of final benchmarking report*.""", normal_style),
                benchmark_fee
            ])

        if data.get('includeIntercompany') == 'on':
            interco_fee = format_currency(data.get('intercompanyAgreementFee', '0')) or '1500'
            tp_data.append([
                Paragraph('<font color="#C00000" face="Roboto-Bold"><b>Inter-company agreement</b></font>', normal_style),
                'One-time',
                Paragraph("""Drafting and finalizing of Inter-company service agreement covering detailed description of service to be provided, components to be included while calculating cost of services, Invoicing period, Receivable cycle, withholding, ownership rights, effective date of agreement, indemnity etc. in compliance with the Transfer Pricing regulations defined under Income tax laws and other applicable Indian laws""", normal_style),
                interco_fee
            ])

        tp_total = 0
        if benchmark_fee is not None:
            try:
                tp_total += int(benchmark_fee.replace(',', ''))
            except:
                pass
        if interco_fee is not None:
            try:
                tp_total += int(interco_fee.replace(',', ''))
            except:
                pass
            total_label_style = ParagraphStyle(
    'TotalLabelStyle',
    parent=small_style,
    alignment=TA_CENTER,
)

        if len(tp_data) > 1:
            tp_data.append([
                Paragraph('<b>Total one-time costs</b>', total_label_style),
                '', '', f'{tp_total:,}'
            ])
            tp_table = Table(tp_data, colWidths=[1.4*inch, 1.1*inch, 3.3*inch, 1.4*inch])
            table_style_tp = [
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor("#FFFFFF")),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.black),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (3, 0), 10),
                ('FONTSIZE', (0, 1), (-1, -1), 8),
                ('ALIGN', (2, 1), (2, -1), 'LEFT'),
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ('GRID', (0, 1), (-1, -1), 0.5, colors.black),
                ('TOPPADDING', (0, 0), (-1, -1), 8),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
                ('LEFTPADDING', (0, 1), (0, -1), 2),
                ('RIGHTPADDING', (0, 0), (-1, -1), 8),
                ('BOX', (0, 0), (-1, -1), 0.5, colors.black),
                ('LINEBELOW', (0, 0), (-1, 0), 0.75, colors.black),
                ('LINEBEFORE', (0, 0), (-1, 0), 0, colors.white),
                ('LINEAFTER', (0, 0), (-1, 0), 0, colors.white),
                ('LINEABOVE', (0, 0), (-1, 0), 0, colors.white),
            ]
            if len(tp_data) > 2:
                total_row_index = len(tp_data) - 1
                table_style_tp.append(('SPAN', (0, total_row_index), (2, total_row_index)))
                table_style_tp.append(('BACKGROUND', (0, total_row_index), (-1, total_row_index), colors.HexColor("#FFFFFF")))
                table_style_tp.append(('ALIGN', (0, total_row_index), (2, total_row_index), 'CENTER'))
            tp_table.setStyle(TableStyle(table_style_tp))
            elements.append(tp_table)
            elements.append(Spacer(1, 10))
            elements.append(Paragraph("*Please note that the above benchmarking report will not be transfer pricing documentation as required to be maintained under transfer pricing regulations. InCorp's empanelled audit partners can assist with the transfer pricing reporting & audit (applicable for companies having intercompany transactions). The quotes for the same can be provided separately.", italic_style))

        elements.append(Spacer(1, 10))
        tp_notes = """<b><u><font color="#002060">Note:</font></u></b><br/><br/>
• All fees quoted above exclude 18% GST.<br/>
• Professional fees exclude all out-of-pocket expenses.<br/>
• Advance of 100% of the above selected option.<br/>"""
        elements.append(Paragraph(tp_notes, normal_style))
        elements.append(Spacer(1, 4))
        elements.append(Paragraph("<b>For Partner: USD 300 per Hour</b>", italic_style))
        elements.append(Spacer(1, 2))
        elements.append(Paragraph("<b>For Associates: USD 200 per Hour</b>", italic_style))
        elements.append(Spacer(1, 6))
        elements.append(Paragraph("^ InCorp's empanelled audit partners can assist with the transfer pricing reporting &amp; audit (applicable for companies having intercompany transactions). The quotes for the same can be provided separately.", italic_style))

    return elements


# ==================== SHARED PDF BUFFER BUILDER ====================

def build_merged_pdf_buffer(data):
    """Build complete merged PDF and return as BytesIO. Used by both routes."""
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer, pagesize=letter,
        rightMargin=0.5*inch, leftMargin=0.5*inch,
        topMargin=1.2*inch, bottomMargin=1*inch
    )
    elements = build_pdf_elements(data)
    doc.build(elements, canvasmaker=InCorpCanvas)

    buffer.seek(0)
    dynamic_pdf = PdfReader(buffer)
    num_dynamic_pages = len(dynamic_pdf.pages)
    print(f"✅ Generated {num_dynamic_pages} dynamic pages")

    merger = PdfWriter()
    print(" → Adding Page 1 (Cover)")
    merger.append(dynamic_pdf, pages=(0, 1))

    static_2_3_4_path = os.path.join(BASE_DIR, 'static_pdfs', 'static_pages_2_3_4.pdf')
    if os.path.exists(static_2_3_4_path):
        print(" → Adding Pages 2-4 (Static)")
        merger.append(static_2_3_4_path)
    else:
        print(f"⚠️ WARNING: {static_2_3_4_path} not found!")

    print(" → Adding Pages 5-13 (Dynamic fee tables)")
    if num_dynamic_pages > 1:
        merger.append(dynamic_pdf, pages=(1, num_dynamic_pages))

    static_14_21_path = os.path.join(BASE_DIR, 'static_pdfs', 'static_pages_14_21.pdf')
    if os.path.exists(static_14_21_path):
        print(" → Adding Pages 14-21 (Static)")
        merger.append(static_14_21_path)
    else:
        print(f"⚠️ WARNING: {static_14_21_path} not found!")

    final_buffer = io.BytesIO()
    merger.write(final_buffer)
    final_buffer.seek(0)
    return final_buffer


# ==================== ROUTES ====================
# FIX 3: Only ONE @app.route('/') above (kept the original one at top of file)

@app.route('/generate_proposal', methods=['POST'])
def generate_proposal():
    """Generate and return the merged PDF proposal."""
    try:
        data = request.json
        final_buffer = build_merged_pdf_buffer(data)
        return send_file(
            final_buffer,
            mimetype='application/pdf',
            as_attachment=True,
            download_name=f'InCorp_Proposal_{data.get("clientCompany", "Client").replace(" ", "_")}_{datetime.now().strftime("%Y%m%d")}.pdf'
        )
    except Exception as e:
        print(f"Error generating PDF: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500


@app.route('/generate_proposal_word', methods=['POST'])
def generate_proposal_word():
    """Generate ONLY dynamic PDF pages and convert to Word - skips static annexure pages."""
    timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
    pdf_temp = os.path.join(BASE_DIR, f'temp_proposal_{timestamp}.pdf')
    client_name = request.json.get("clientCompany", "Client").replace(" ", "_") if request.json else "Client"
    docx_file = os.path.join(BASE_DIR, f'InCorp_Proposal_{client_name}_{datetime.now().strftime("%Y%m%d")}.docx')

    try:
        data = request.json

        # ── Step 1: Build full merged PDF (same as PDF download) ──
        print("📄 Building merged PDF...")
        final_buffer = build_merged_pdf_buffer(data)
        with open(pdf_temp, 'wb') as f:
            f.write(final_buffer.read())
        print(f"✅ Merged PDF saved: {pdf_temp}")

        # ── Step 2: PDF → Images using pymupdf (pip install pymupdf) ──
        try:
            import fitz  # pymupdf
        except ImportError:
            raise Exception("pymupdf not installed. Run: pip install pymupdf")

        from docx import Document as DocxDocument
        from docx.shared import Inches, Pt
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        print("🖼️ Rendering PDF pages to images...")
        pdf_doc = fitz.open(pdf_temp)
        img_files = []
        mat = fitz.Matrix(2.0, 2.0)  # 144 DPI - good quality

        for page_num in range(len(pdf_doc)):
            page = pdf_doc[page_num]
            pix = page.get_pixmap(matrix=mat)
            img_path = pdf_temp.replace('.pdf', f'_pg{page_num+1:03d}.png')
            pix.save(img_path)
            img_files.append(img_path)

        pdf_doc.close()
        print(f"✅ {len(img_files)} pages rendered")

        # ── Step 3: Images → Word (one full-page image per page) ──
        word_doc = DocxDocument()

        # Set page size/margins BEFORE clearing body
        section = word_doc.sections[0]
        section.page_width    = int(8.5 * 914400)
        section.page_height   = int(11  * 914400)
        section.left_margin   = int(0.0 * 914400)
        section.right_margin  = int(0.0 * 914400)
        section.top_margin    = int(0.0 * 914400)
        section.bottom_margin = int(0.0 * 914400)

        # Now clear default content (paragraph only, keep sectPr)
        from docx.oxml.ns import qn as _qn
        for el in list(word_doc.element.body):
            if not el.tag.endswith('}sectPr'):
                word_doc.element.body.remove(el)

        for i, img_path in enumerate(img_files):
            p = word_doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(0)
            run = p.add_run()
            run.add_picture(img_path, width=Inches(8.5))

            if i < len(img_files) - 1:
                br = OxmlElement('w:br')
                br.set(qn('w:type'), 'page')
                run._r.append(br)

        word_doc.save(docx_file)
        print(f"✅ Word saved: {docx_file}")

        # Cleanup
        for f in [pdf_temp] + img_files:
            try:
                if os.path.exists(f): os.remove(f)
            except: pass

        return send_file(
            docx_file,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=os.path.basename(docx_file)
        )

    except Exception as e:
        print(f"❌ Error generating Word: {str(e)}")
        import traceback
        traceback.print_exc()
        for f in [pdf_temp, docx_file]:
            try:
                if os.path.exists(f): os.remove(f)
            except: pass
        return jsonify({'error': str(e)}), 500
    

if __name__ == '__main__':
    os.makedirs('static_pdfs', exist_ok=True)
    
    
    print("="*60)
    print("InCorp Proposal Generator - FINAL VERSION")
    print("="*60)
    print("\n✅ 3 Default accounting rows (Upto 20, 20-50, 50-80)")
    print("✅ 2 Default payroll rows (Upto 10, 11-20)")
    print("✅ Monthly fee fields added for nested tables")
    print("✅ Cover page with company name at 1.8 inch")
    print("✅ Header/Footer removed from Page 1")
    print("✅ Total costs calculated automatically")
    print("\n🚀 Starting server on http://localhost:5000")
    print("="*60)
    app.run(debug=True, host='0.0.0.0', port=5000)